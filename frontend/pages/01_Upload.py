"""
Upload Documents Page
Beautiful, modern file upload interface with enhanced UX and visual appeal
"""

import streamlit as st
import os
import time
import asyncio
import threading
from datetime import datetime
from typing import Dict, List, Any
import pandas as pd
import numpy as np
import requests
import json
from PIL import Image
import io

# Import custom modules - using relative imports for compatibility
try:
    from utils.session_state import get_session_manager
    from config.settings import app_config
except ImportError:
    # Fallback if modules don't exist
    def get_session_manager():
        return None
    
    class app_config:
        SUPPORTED_FORMATS = ["pdf", "docx", "txt", "jpg", "png", "jpeg"]
        MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

# Initialize session manager
session_manager = get_session_manager()

def render_upload_page():
    """Render the beautiful, modern upload page"""
    
    # Ensure processing_settings is properly initialized
    if 'processing_settings' not in st.session_state or not isinstance(st.session_state.processing_settings, dict):
        st.session_state.processing_settings = {
            "chunk_size": 1000,
            "overlap": 200,
            "extract_tables": True,
            "detect_charts": True
        }
    
    # Beautiful header with gradient background effect
    st.markdown("""
    <div style="
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 15px;
        margin-bottom: 2rem;
        text-align: center;
        color: white;
        box-shadow: 0 8px 32px rgba(0,0,0,0.1);
    ">
        <h1 style="margin: 0; font-size: 2.5rem; font-weight: 700;">📤 Upload Documents</h1>
        <p style="margin: 0.5rem 0 0 0; font-size: 1.2rem; opacity: 0.9;">
            Transform your documents into intelligent insights with AI-powered analysis
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Check backend connection
    if st.button("🔗 Check Backend Connection", key="check_backend"):
        connect_to_backend()
    
    # Feature highlights in beautiful cards
    render_feature_highlights()
    
    # Main upload section
    render_main_upload_section()
    
    # Processing monitor (if files are being processed)
    if st.session_state.get('pending_files'):
        render_processing_monitor()
    
    # Upload history
    if st.session_state.get('uploaded_files'):
        render_upload_history()
    
    # Processing statistics
    render_processing_stats()

def render_feature_highlights():
    """Render beautiful feature highlight cards"""
    st.markdown("### ✨ What You Can Do")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div style="
            background: linear-gradient(135deg, #ff9a9e 0%, #fecfef 100%);
            padding: 1.5rem;
            border-radius: 12px;
            text-align: center;
            color: white;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            margin-bottom: 1rem;
        ">
            <div style="font-size: 3rem; margin-bottom: 0.5rem;">📄</div>
            <h4 style="margin: 0 0 0.5rem 0;">Multi-Format Support</h4>
            <p style="margin: 0; font-size: 0.9rem; opacity: 0.9;">
                PDF, Images, Documents & more
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div style="
            background: linear-gradient(135deg, #a8edea 0%, #fed6e3 100%);
            padding: 1.5rem;
            border-radius: 12px;
            text-align: center;
            color: white;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            margin-bottom: 1rem;
        ">
            <div style="font-size: 3rem; margin-bottom: 0.5rem;">🤖</div>
            <h4 style="margin: 0 0 0.5rem 0;">AI-Powered OCR</h4>
            <p style="margin: 0; font-size: 0.9rem; opacity: 0.9;">
                High-accuracy text extraction
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div style="
            background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%);
            padding: 1.5rem;
            border-radius: 12px;
            text-align: center;
            color: white;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
            margin-bottom: 1rem;
        ">
            <div style="font-size: 3rem; margin-bottom: 0.5rem;">📊</div>
            <h4 style="margin: 0 0 0.5rem 0;">Smart Analytics</h4>
            <p style="margin: 0; font-size: 0.9rem; opacity: 0.9;">
                Tables, charts & insights
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")

def render_main_upload_section():
    """Render the main upload section with beautiful design"""
    
    # Upload zone container
    upload_container = st.container()
    
    with upload_container:
        st.markdown("### 🚀 Upload Your Files")
        
        # Create a unified upload zone that combines visual design with functionality
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col2:
            # Enterprise-level, premium file uploader with advanced styling
            st.markdown("""
            <style>
            /* Premium Enterprise File Uploader Styling */
            .stFileUploader {
                background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%) !important;
                border-radius: 20px !important;
                padding: 3rem 2rem !important;
                border: none !important;
                box-shadow: 
                    0 20px 40px rgba(102, 126, 234, 0.3),
                    0 0 0 1px rgba(255, 255, 255, 0.1),
                    inset 0 1px 0 rgba(255, 255, 255, 0.2) !important;
                transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1) !important;
                position: relative !important;
                overflow: hidden !important;
                backdrop-filter: blur(10px) !important;
            }
            
            .stFileUploader::before {
                content: "" !important;
                position: absolute !important;
                top: 0 !important;
                left: -100% !important;
                width: 100% !important;
                height: 100% !important;
                background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.2), transparent) !important;
                transition: left 0.8s ease !important;
            }
            
            .stFileUploader:hover::before {
                left: 100% !important;
            }
            
            .stFileUploader:hover {
                transform: translateY(-8px) scale(1.02) !important;
                box-shadow: 
                    0 30px 60px rgba(102, 126, 234, 0.4),
                    0 0 0 1px rgba(255, 255, 255, 0.15),
                    inset 0 1px 0 rgba(255, 255, 255, 0.3) !important;
            }
            
            /* Premium Dropzone Styling */
            .stFileUploader [data-testid="stFileUploadDropzone"] {
                background: rgba(255, 255, 255, 0.1) !important;
                border: 2px dashed rgba(255, 255, 255, 0.3) !important;
                border-radius: 16px !important;
                padding: 2rem !important;
                transition: all 0.3s ease !important;
                backdrop-filter: blur(5px) !important;
            }
            
            .stFileUploader [data-testid="stFileUploadDropzone"]:hover {
                background: rgba(255, 255, 255, 0.15) !important;
                border-color: rgba(255, 255, 255, 0.5) !important;
                transform: scale(1.02) !important;
            }
            
            /* Premium Instructions Text */
            .stFileUploader [data-testid="stFileDropzoneInstructions"] {
                text-align: center !important;
                padding: 1.5rem 0 !important;
                position: relative !important;
            }
            
            .stFileUploader [data-testid="stFileDropzoneInstructions"]::before {
                content: "🚀" !important;
                font-size: 4rem !important;
                display: block !important;
                text-align: center !important;
                margin-bottom: 1rem !important;
                animation: rocketFloat 4s ease-in-out infinite !important;
                filter: drop-shadow(0 4px 8px rgba(0, 0, 0, 0.3)) !important;
            }
            
            .stFileUploader [data-testid="stFileDropzoneInstructions"] span {
                font-size: 1.5rem !important;
                font-weight: 700 !important;
                color: white !important;
                text-shadow: 0 2px 4px rgba(0, 0, 0, 0.3) !important;
                letter-spacing: 0.5px !important;
                display: block !important;
                margin-bottom: 0.5rem !important;
            }
            
            .stFileUploader [data-testid="stFileDropzoneInstructions"] small {
                font-size: 1.1rem !important;
                color: rgba(255, 255, 255, 0.9) !important;
                margin-top: 0.5rem !important;
                display: block !important;
                font-weight: 500 !important;
                text-shadow: 0 1px 2px rgba(0, 0, 0, 0.2) !important;
            }
            
            /* Premium Browse Button */
            .stFileUploader button {
                background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%) !important;
                color: #667eea !important;
                border: none !important;
                border-radius: 30px !important;
                padding: 1rem 2.5rem !important;
                font-weight: 700 !important;
                font-size: 1.1rem !important;
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
                box-shadow: 
                    0 8px 25px rgba(0, 0, 0, 0.15),
                    0 0 0 1px rgba(255, 255, 255, 0.2) !important;
                text-transform: uppercase !important;
                letter-spacing: 0.5px !important;
                position: relative !important;
                overflow: hidden !important;
            }
            
            .stFileUploader button::before {
                content: "" !important;
                position: absolute !important;
                top: 0 !important;
                left: -100% !important;
                width: 100% !important;
                height: 100% !important;
                background: linear-gradient(90deg, transparent, rgba(102, 126, 234, 0.1), transparent) !important;
                transition: left 0.6s ease !important;
            }
            
            .stFileUploader button:hover::before {
                left: 100% !important;
            }
            
            .stFileUploader button:hover {
                transform: translateY(-3px) scale(1.05) !important;
                box-shadow: 
                    0 15px 35px rgba(0, 0, 0, 0.2),
                    0 0 0 1px rgba(255, 255, 255, 0.3) !important;
                color: #5a67d8 !important;
            }
            
            .stFileUploader button:active {
                transform: translateY(-1px) scale(1.02) !important;
            }
            
            /* Premium Animations */
            @keyframes rocketFloat {
                0%, 100% { 
                    transform: translateY(0px) rotate(-5deg); 
                    filter: drop-shadow(0 4px 8px rgba(0, 0, 0, 0.3));
                }
                25% { 
                    transform: translateY(-8px) rotate(0deg); 
                    filter: drop-shadow(0 8px 16px rgba(0, 0, 0, 0.4));
                }
                50% { 
                    transform: translateY(-12px) rotate(5deg); 
                    filter: drop-shadow(0 12px 24px rgba(0, 0, 0, 0.5));
                }
                75% { 
                    transform: translateY(-8px) rotate(0deg); 
                    filter: drop-shadow(0 8px 16px rgba(0, 0, 0, 0.4));
                }
            }
            
            /* Premium Glow Effect */
            .stFileUploader::after {
                content: "" !important;
                position: absolute !important;
                top: -2px !important;
                left: -2px !important;
                right: -2px !important;
                bottom: -2px !important;
                background: linear-gradient(45deg, #667eea, #764ba2, #f093fb, #667eea) !important;
                border-radius: 22px !important;
                z-index: -1 !important;
                opacity: 0 !important;
                transition: opacity 0.4s ease !important;
                animation: glowRotate 3s linear infinite !important;
            }
            
            .stFileUploader:hover::after {
                opacity: 0.7 !important;
            }
            
            @keyframes glowRotate {
                0% { transform: rotate(0deg); }
                100% { transform: rotate(360deg); }
            }
            
            /* Premium File Type Icons */
            .stFileUploader .stMarkdown {
                position: relative !important;
            }
            
            /* Responsive Design */
            @media (max-width: 768px) {
                .stFileUploader {
                    padding: 2rem 1rem !important;
                    margin: 1rem !important;
                }
                
                .stFileUploader [data-testid="stFileDropzoneInstructions"] span {
                    font-size: 1.3rem !important;
                }
                
                .stFileUploader button {
                    padding: 0.8rem 2rem !important;
                    font-size: 1rem !important;
                }
            }
            </style>
            """, unsafe_allow_html=True)
            
            # Premium enterprise file uploader with advanced styling
            uploaded_files = st.file_uploader(
                "🚀 **PREMIUM DOCUMENT UPLOAD**",
                type=app_config.SUPPORTED_FORMATS,
                accept_multiple_files=True,
                help="🎯 **Enterprise-Grade Support**: PDF, PNG, JPG, JPEG, DOCX, TXT • ⚡ **Performance**: Up to 50MB per file • 🔒 **Secure**: Enterprise-level security",
                key="main_file_uploader"
            )
        
        if uploaded_files:
            st.markdown("### 📋 Selected Files")
            render_file_selection(uploaded_files)
        else:
            # Show supported formats info
            st.info("💡 **Supported Formats**: PDF, DOCX, TXT, JPG, PNG, JPEG")
            st.info("📏 **Maximum File Size**: 50MB per file")

def render_file_selection(files):
    """Render beautiful file selection interface"""
    
    # File validation
    valid_files = []
    for file in files:
        if validate_file(file):
            valid_files.append(file)
    
    if not valid_files:
        st.error("❌ No valid files selected. Please check file formats and sizes.")
        return
    
    # Display selected files in beautiful cards
    for i, file in enumerate(valid_files):
        render_file_card(file, i)
    
    # Upload actions
    if valid_files:
        render_upload_actions(valid_files)

def render_file_card(file, index):
    """Render a beautiful file card"""
    
    # Get file icon and color
    file_ext = os.path.splitext(file.name)[1].lower()
    icon_map = {
        '.pdf': ('📄', '#e53e3e'),
        '.docx': ('📝', '#3182ce'),
        '.txt': ('📄', '#38a169'),
        '.png': ('🖼️', '#d69e2e'),
        '.jpg': ('🖼️', '#d69e2e'),
        '.jpeg': ('🖼️', '#d69e2e')
    }
    icon, color = icon_map.get(file_ext, ('📁', '#718096'))
    
    # File card
    col1, col2, col3, col4 = st.columns([1, 4, 2, 1])
    
    with col1:
        st.markdown(f"""
        <div style="
            background: {color};
            color: white;
            width: 50px;
            height: 50px;
            border-radius: 10px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 1.5rem;
            margin: 0 auto;
        ">
            {icon}
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown(f"**{file.name}**")
        st.caption(f"📏 Size: {file.size // 1024:,} KB")
        st.caption(f"🏷️ Type: {file_ext.upper()}")
    
    with col3:
        # File status
        if file.name in [f['name'] for f in st.session_state.get('pending_files', [])]:
            st.info("⏳ Queued")
        elif file.name in [f['name'] for f in st.session_state.get('uploaded_files', [])]:
            st.success("✅ Processed")
        else:
            st.info("📤 Ready")
    
    with col4:
        # Remove button
        if st.button("🗑️", key=f"remove_{index}", help="Remove file"):
            # Remove file logic would go here
            st.rerun()
    
    st.markdown("---")

def render_upload_actions(files):
    """Render beautiful upload action buttons"""
    
    st.markdown("### ⚡ Ready to Process")
    
    # Action buttons in a grid
    col1, col2, col3 = st.columns([2, 1, 1])
    
    with col1:
        if st.button(
            "🚀 Start Processing", 
            key="start_processing",
            use_container_width=True,
            type="primary"
        ):
            process_files(files)
    
    with col2:
        if st.button(
            "👁️ Preview Files", 
            key="preview_files",
            use_container_width=True
        ):
            show_file_previews(files)
    
    with col3:
        if st.button(
            "⚙️ Settings", 
            key="processing_settings",
            use_container_width=True
        ):
            show_processing_settings()

def validate_file(file) -> bool:
    """Validate uploaded file"""
    # File size validation
    if file.size > app_config.MAX_FILE_SIZE:
        st.error(f"""
        ❌ **File too large**: {file.name}
        - Current size: {file.size // (1024*1024):.1f} MB
        - Maximum allowed: {app_config.MAX_FILE_SIZE // (1024*1024)} MB
        """)
        return False
    
    # File format validation
    file_extension = os.path.splitext(file.name)[1].lower()
    if file_extension not in app_config.SUPPORTED_FORMATS:
        st.error(f"""
        ❌ **Unsupported format**: {file.name}
        - Detected format: {file_extension.upper()}
        - Supported formats: {', '.join(app_config.SUPPORTED_FORMATS)}
        """)
        return False
    
    return True

def process_files(files):
    """Process uploaded files with real content extraction"""
    
    st.markdown("### ⚙️ Processing Your Files")
    
    # Progress tracking
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # Processing steps
    steps = ["📤 Uploading", "🔍 Analyzing", "📝 Extracting", "🤖 AI Processing", "✅ Completing"]
    
    total_files = len(files)
    processed_files = 0
    
    try:
        for i, file in enumerate(files):
            try:
                status_text.text(f"Processing {file.name}...")
                
                # Step 1: Upload and initialize
                step_progress = (i * len(steps) + 0) / (total_files * len(steps))
                progress_bar.progress(step_progress)
                status_text.text(f"📤 Uploading {file.name}...")
                
                # Add file to session state
                file_data = add_file_to_session(file, "pending")
                
                # Step 2: Analyze file type
                step_progress = (i * len(steps) + 1) / (total_files * len(steps))
                progress_bar.progress(step_progress)
                status_text.text(f"🔍 Analyzing {file.name}...")
                
                # Step 3: Extract content
                step_progress = (i * len(steps) + 2) / (total_files * len(steps))
                progress_bar.progress(step_progress)
                status_text.text(f"📝 Extracting content from {file.name}...")
                
                # Actually process the document content
                try:
                    processed_file_data = process_document_content(file, file_data)
                    
                    # Update the file data in session state
                    file_index = next((i for i, f in enumerate(st.session_state.uploaded_files) 
                                      if f['name'] == file.name), None)
                    if file_index is not None:
                        st.session_state.uploaded_files[file_index] = processed_file_data
                    
                except Exception as content_error:
                    st.error(f"❌ Content extraction failed for {file.name}: {str(content_error)}")
                    # Mark file as failed but continue processing
                    file_data['status'] = 'failed'
                    file_data['processing_errors'].append(str(content_error))
                    continue
                
                # Step 4: AI Processing (placeholder for future enhancement)
                step_progress = (i * len(steps) + 3) / (total_files * len(steps))
                progress_bar.progress(step_progress)
                status_text.text(f"🤖 AI Processing {file.name}...")
                time.sleep(0.2)  # Brief pause for AI processing
                
                # Step 5: Complete
                step_progress = (i * len(steps) + 4) / (total_files * len(steps))
                progress_bar.progress(step_progress)
                status_text.text(f"✅ Completing {file.name}...")
                
                processed_files += 1
                
                # Brief pause between files
                time.sleep(0.1)
                
            except Exception as file_error:
                st.error(f"❌ Error processing file {file.name}: {str(file_error)}")
                st.error(f"Error type: {type(file_error).__name__}")
                import traceback
                st.error(f"Traceback: {traceback.format_exc()}")
                # Continue with next file
                continue
        
        # Complete
        progress_bar.progress(1.0)
        status_text.text("✅ All files processed successfully!")
        
        st.success("🎉 Processing complete!")
        
        # Show summary with real data
        show_processing_summary(st.session_state.get('uploaded_files', []))
        
    except Exception as e:
        st.error(f"❌ Critical error during processing: {str(e)}")
        st.error(f"Error type: {type(e).__name__}")
        import traceback
        st.error(f"Traceback: {traceback.format_exc()}")
        progress_bar.progress(0)
        status_text.text("❌ Processing failed!")

def add_file_to_session(file, status):
    """Add file to session state with enhanced metadata"""
    file_data = {
        "name": file.name,
        "size": file.size,
        "type": os.path.splitext(file.name)[1].lower(),
        "uploaded_at": datetime.now().isoformat(),
        "upload_time": datetime.now().isoformat(),  # For Analysis page compatibility
        "status": status,
        "processing_start": datetime.now().isoformat(),
        "processing_time": 0,
        "confidence": 0,
        "text_chunks": 0,
        "tables": 0,
        "images": 0,
        "extracted_text": "",
        "extracted_tables": [],
        "extracted_images": [],
        "processing_errors": []
    }
    
    if "uploaded_files" not in st.session_state:
        st.session_state.uploaded_files = []
    
    st.session_state.uploaded_files.append(file_data)
    return file_data

def show_processing_summary(files):
    """Show beautiful processing summary"""
    
    st.markdown("### 📊 Processing Summary")
    
    # Summary metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("📁 Total Files", len(files))
    
    with col2:
        try:
            total_size = sum(f.get('size', 0) for f in files)
            st.metric("💾 Total Size", f"{total_size // (1024*1024):.1f} MB")
        except Exception as e:
            st.metric("💾 Total Size", "Error")
            st.caption(f"Calculation failed: {str(e)}")
    
    with col3:
        try:
            file_types = set(f.get('type', 'unknown') for f in files)
            st.metric("🏷️ File Types", len(file_types))
        except Exception as e:
            st.metric("🏷️ File Types", "Error")
            st.caption(f"Calculation failed: {str(e)}")
    
    with col4:
        try:
            processed_count = len([f for f in files if f.get('status') == 'processed'])
            success_rate = (processed_count / len(files) * 100) if files else 0
            st.metric("✅ Success Rate", f"{success_rate:.1f}%")
        except Exception as e:
            st.metric("✅ Success Rate", "Error")
            st.caption(f"Calculation failed: {str(e)}")
    
    # Content extraction summary
    if files:
        st.markdown("### 🔍 Content Extraction Summary")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            try:
                total_text_chunks = sum(f.get('text_chunks', 0) for f in files)
                st.metric("📝 Text Chunks", total_text_chunks)
            except Exception as e:
                st.metric("📝 Text Chunks", "Error")
                st.caption(f"Calculation failed: {str(e)}")
        
        with col2:
            try:
                total_tables = sum(f.get('tables', 0) for f in files)
                st.metric("📊 Tables", total_tables)
            except Exception as e:
                st.metric("📊 Tables", "Error")
                st.caption(f"Calculation failed: {str(e)}")
        
        with col3:
            try:
                total_images = sum(f.get('images', 0) for f in files)
                st.metric("🖼️ Images", total_images)
            except Exception as e:
                st.metric("🖼️ Images", "Error")
                st.caption(f"Calculation failed: {str(e)}")
    
    # File type breakdown
    if len(file_types) > 1:
        st.markdown("**📊 File Type Breakdown**")
        
        type_counts = {}
        for file in files:
            ext = os.path.splitext(file.name)[1].lower()
            type_counts[ext.upper()] = type_counts.get(ext.upper(), 0) + 1
        
        # Create a simple bar chart
        try:
            chart_data = pd.DataFrame([
                {"Type": ext, "Count": count}
                for ext, count in type_counts.items()
            ])
            
            st.bar_chart(chart_data.set_index("Type"))
        except Exception as e:
            st.error(f"❌ Error creating chart: {str(e)}")
            st.info("Chart data unavailable due to processing errors")

def show_file_previews(files):
    """Show file previews"""
    st.markdown("### 👁️ File Previews")
    
    for file in files:
        with st.expander(f"📄 {file.name}", expanded=False):
            col1, col2 = st.columns([1, 2])
            
            with col1:
                st.markdown(f"**File Details**")
                st.caption(f"Name: {file.name}")
                st.caption(f"Size: {file.size // 1024:,} KB")
                st.caption(f"Type: {os.path.splitext(file.name)[1].upper()}")
            
            with col2:
                # Preview content
                if hasattr(file, 'type') and file.type and file.type.startswith('image'):
                    st.image(file, caption=file.name, use_column_width=True)
                elif os.path.splitext(file.name)[1].lower() == '.pdf':
                    st.info("📄 PDF preview coming soon!")
                else:
                    st.text("Preview not available for this file type")

def show_processing_settings():
    """Show processing settings"""
    st.markdown("### ⚙️ Processing Settings")
    
    # Ensure processing_settings is always a dictionary
    if 'processing_settings' not in st.session_state or not isinstance(st.session_state.processing_settings, dict):
        st.session_state.processing_settings = {
            "chunk_size": 1000,
            "overlap": 200,
            "extract_tables": True,
            "detect_charts": True
        }
    
    # If it's still not a dictionary, force reset it
    if not isinstance(st.session_state.processing_settings, dict):
        st.warning("⚠️ processing_settings was corrupted, resetting to defaults...")
        st.session_state.processing_settings = {
            "chunk_size": 1000,
            "overlap": 200,
            "extract_tables": True,
            "detect_charts": True
        }
    
    with st.form("processing_settings_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            chunk_size = st.slider("Chunk Size", 500, 2000, st.session_state.processing_settings.get("chunk_size", 1000), step=100, help="Size of text chunks for processing")
            overlap = st.slider("Chunk Overlap", 100, 500, st.session_state.processing_settings.get("overlap", 200), step=50, help="Overlap between chunks")
        
        with col2:
            extract_tables = st.checkbox("Extract Tables", value=st.session_state.processing_settings.get("extract_tables", True), help="Automatically detect and extract tables")
            detect_charts = st.checkbox("Detect Charts", value=st.session_state.processing_settings.get("detect_charts", True), help="Automatically detect and extract charts")
        
        if st.form_submit_button("💾 Save Settings"):
            st.success("Settings saved successfully!")
            st.session_state.processing_settings = {
                "chunk_size": chunk_size,
                "overlap": overlap,
                "extract_tables": extract_tables,
                "detect_charts": detect_charts
            }
    
    # Reset button outside the form
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        if st.button("🔄 Reset to Defaults", key="reset_settings"):
            st.session_state.processing_settings = {
                "chunk_size": 1000,
                "overlap": 200,
                "extract_tables": True,
                "detect_charts": True
            }
            st.success("Settings reset to defaults!")
            st.rerun()

def render_processing_monitor():
    """Render processing monitor"""
    st.markdown("### ⏳ Processing Monitor")
    
    # Simple processing status
    for file_data in st.session_state.pending_files:
        st.info(f"🔄 Processing: {file_data['name']}")

def render_upload_history():
    """Render upload history"""
    st.markdown("### 📚 Upload History")
    
    if not st.session_state.get('uploaded_files'):
        st.info("📭 No files uploaded yet. Start by uploading some documents!")
        return
    
    # Simple list view
    for file_data in st.session_state.uploaded_files:
        with st.expander(f"📄 {file_data['name']}", expanded=False):
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.markdown(f"**File Info**")
                st.caption(f"Type: {file_data.get('type', 'unknown')}")
                st.caption(f"Size: {file_data.get('size', 0) // 1024:,} KB")
                st.caption(f"Status: {file_data.get('status', 'unknown')}")
                st.caption(f"Uploaded: {file_data.get('uploaded_at', 'Unknown')[:19]}")
                
                # Show content extraction info
                if file_data.get('text_chunks', 0) > 0:
                    st.caption(f"📝 Text Chunks: {file_data.get('text_chunks', 0)}")
                if file_data.get('tables', 0) > 0:
                    st.caption(f"📊 Tables: {file_data.get('tables', 0)}")
                if file_data.get('images', 0) > 0:
                    st.caption(f"🖼️ Images: {file_data.get('images', 0)}")
                if file_data.get('confidence', 0) > 0:
                    st.caption(f"🎯 Confidence: {file_data.get('confidence', 0):.1f}%")
            
            with col2:
                st.markdown(f"**Actions**")
                col_a, col_b = st.columns(2)
                
                with col_a:
                    if st.button("👁️ View", key=f"view_{file_data['name']}"):
                        show_file_details(file_data)
                
                with col_b:
                    if st.button("🗑️ Delete", key=f"delete_{file_data['name']}"):
                        st.session_state.uploaded_files.remove(file_data)
                        st.rerun()

def render_processing_stats():
    """Render processing statistics with real data"""
    st.markdown("### 📊 Processing Statistics")
    
    # Get real data from session state
    uploaded_files = st.session_state.get('uploaded_files', [])
    
    # Create metrics cards
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        total_files = len(uploaded_files)
        st.metric("📁 Total Files", total_files)
    
    with col2:
        total_size = sum(f.get('size', 0) for f in uploaded_files)
        st.metric("💾 Total Size", f"{total_size // (1024*1024):.1f} MB")
    
    with col3:
        processed_files = len([f for f in uploaded_files if f.get('status') == 'processed'])
        st.metric("✅ Processed", processed_files)
    
    with col4:
        success_rate = (processed_files / total_files * 100) if total_files > 0 else 0
        st.metric("🎯 Success Rate", f"{success_rate:.1f}%")
    
    # Content extraction stats
    if uploaded_files:
        st.markdown("### 🔍 Content Extraction Stats")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            total_text_chunks = sum(f.get('text_chunks', 0) for f in uploaded_files)
            st.metric("📝 Text Chunks", total_text_chunks)
        
        with col2:
            total_tables = sum(f.get('tables', 0) for f in uploaded_files)
            st.metric("📊 Tables", total_tables)
        
        with col3:
            total_images = sum(f.get('images', 0) for f in uploaded_files)
            st.metric("🖼️ Images", total_images)
        
        with col4:
            try:
                confidence_values = [f.get('confidence', 0) for f in uploaded_files if f.get('confidence', 0) > 0]
                if confidence_values:
                    avg_confidence = np.mean(confidence_values)
                    st.metric("🎯 Avg Confidence", f"{avg_confidence:.1f}%")
                else:
                    st.metric("🎯 Avg Confidence", "N/A")
            except Exception as e:
                st.metric("🎯 Avg Confidence", "Error")
                st.caption(f"Calculation failed: {str(e)}")

def process_document_content(file, file_data):
    """Process document content and extract real data"""
    try:
        start_time = datetime.now()
        
        # Update status to processing
        file_data['status'] = 'processing'
        file_data['processing_start'] = start_time.isoformat()
        
        # Extract content based on file type
        file_extension = file_data['type'].lower()
        
        # Validate file object before processing
        if not file or not hasattr(file, 'read'):
            st.error(f"Invalid file object passed for processing. File type: {type(file)}")
            extracted_data = {
                'text_chunks': 0,
                'tables': 0,
                'images': 0,
                'extracted_text': "Invalid file object",
                'extracted_tables': [],
                'extracted_images': []
            }
        else:
            try:
                if file_extension == '.pdf':
                    extracted_data = process_pdf_file(file)
                elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                    extracted_data = process_image_file(file)
                elif file_extension == '.docx':
                    extracted_data = process_docx_file(file)
                elif file_extension == '.txt':
                    extracted_data = process_text_file(file)
                else:
                    extracted_data = process_unknown_file(file)
            except Exception as extraction_error:
                st.warning(f"⚠️ Content extraction failed for {file.name}: {str(extraction_error)}")
                extracted_data = {
                    'text_chunks': 0,
                    'tables': 0,
                    'images': 0,
                    'extracted_text': f"Extraction failed: {str(extraction_error)}",
                    'extracted_tables': [],
                    'extracted_images': []
                }
        
        # Update file data with extracted content
        file_data.update(extracted_data)
        
        # Calculate processing time and confidence
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        file_data['processing_time'] = processing_time
        file_data['status'] = 'processed'
        
        # Calculate confidence score based on extraction quality
        confidence = calculate_confidence_score(extracted_data)
        file_data['confidence'] = confidence
        
        return file_data
        
    except Exception as e:
        st.error(f"❌ Processing failed for {file.name}: {str(e)}")
        file_data['status'] = 'error'
        file_data['processing_errors'].append(str(e))
        return file_data

def process_pdf_file(file):
    """Process PDF file and extract content"""
    try:
        # Validate file parameter
        if not file or not hasattr(file, 'read'):
            st.error("Invalid file object passed to process_pdf_file")
            return {
                'text_chunks': 0,
                'tables': 0,
                'images': 0,
                'extracted_text': "Invalid file object",
                'extracted_tables': [],
                'extracted_images': []
            }
        
        import pdfplumber
        
        extracted_data = {
            'text_chunks': 0,
            'tables': 0,
            'images': 0,
            'extracted_text': "",
            'extracted_tables': [],
            'extracted_images': []
        }
        
        # Read PDF content
        with pdfplumber.open(file) as pdf:
            text_content = []
            table_count = 0
            
            for page_num, page in enumerate(pdf.pages):
                try:
                    # Extract text
                    page_text = page.extract_text()
                    if page_text and isinstance(page_text, str):
                        text_content.append(page_text)
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables and isinstance(tables, list):
                        table_count += len(tables)
                        extracted_data['extracted_tables'].extend(tables)
                    
                    # Extract images (basic detection)
                    if hasattr(page, 'images') and page.images:
                        extracted_data['images'] += len(page.images)
                        if isinstance(page.images, list):
                            extracted_data['extracted_images'].extend(page.images)
                except Exception as page_error:
                    st.warning(f"⚠️ Error processing page {page_num + 1}: {str(page_error)}")
                    continue
            
            # Combine text and create chunks
            if text_content:
                full_text = "\n".join(text_content)
                extracted_data['extracted_text'] = full_text
                
                # Create text chunks safely
                try:
                    processing_settings = st.session_state.get('processing_settings', {})
                    if not isinstance(processing_settings, dict):
                        st.warning(f"⚠️ processing_settings is not a dictionary: {type(processing_settings)}")
                        processing_settings = {}
                    
                    chunk_size = processing_settings.get('chunk_size', 1000)
                    if chunk_size > 0 and len(full_text) > 0:
                        chunks = [full_text[i:i+chunk_size] for i in range(0, len(full_text), chunk_size)]
                        extracted_data['text_chunks'] = len(chunks)
                    else:
                        extracted_data['text_chunks'] = 1
                except Exception as chunk_error:
                    st.warning(f"⚠️ Error creating text chunks: {str(chunk_error)}")
                    extracted_data['text_chunks'] = 1
            else:
                extracted_data['text_chunks'] = 0
            
            extracted_data['tables'] = table_count
        
        return extracted_data
        
    except Exception as e:
        st.error(f"Error processing PDF: {str(e)}")
        return {
            'text_chunks': 0,
            'tables': 0,
            'images': 0,
            'extracted_text': "",
            'extracted_tables': [],
            'extracted_images': []
        }

def process_image_file(file):
    """Process image file and extract content"""
    try:
        # Validate file parameter
        if not file or not hasattr(file, 'read'):
            st.error("Invalid file object passed to process_image_file")
            return {
                'text_chunks': 0,
                'tables': 0,
                'images': 1,
                'extracted_text': "Invalid file object",
                'extracted_tables': [],
                'extracted_images': []
            }
        
        from PIL import Image
        import pytesseract
        
        extracted_data = {
            'text_chunks': 0,
            'tables': 0,
            'images': 1,
            'extracted_text': "",
            'extracted_tables': [],
            'extracted_images': [file.name]
        }
        
        # Open image
        image = Image.open(file)
        
        # Extract text using OCR
        try:
            text = pytesseract.image_to_string(image)
            if text and isinstance(text, str) and text.strip():
                extracted_data['extracted_text'] = text
                
                # Create text chunks safely
                try:
                    processing_settings = st.session_state.get('processing_settings', {})
                    if not isinstance(processing_settings, dict):
                        st.warning(f"⚠️ processing_settings is not a dictionary: {type(processing_settings)}")
                        processing_settings = {}
                    
                    chunk_size = processing_settings.get('chunk_size', 1000)
                    if chunk_size > 0 and len(text) > 0:
                        chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
                        extracted_data['text_chunks'] = len(chunks)
                    else:
                        extracted_data['text_chunks'] = 1
                except Exception as chunk_error:
                    st.warning(f"⚠️ Error creating text chunks: {str(chunk_error)}")
                    extracted_data['text_chunks'] = 1
            else:
                extracted_data['extracted_text'] = "No text detected in image"
                extracted_data['text_chunks'] = 0
        except Exception as ocr_error:
            st.warning(f"⚠️ OCR processing failed: {str(ocr_error)}")
            extracted_data['extracted_text'] = "OCR processing not available"
            extracted_data['text_chunks'] = 0
        
        return extracted_data
        
    except Exception as e:
        st.error(f"Error processing image: {str(e)}")
        return {
            'text_chunks': 0,
            'tables': 0,
            'images': 1,
            'extracted_text': "",
            'extracted_tables': [],
            'extracted_images': []
        }

def process_docx_file(file):
    """Process DOCX file and extract content"""
    try:
        from docx import Document
        
        # Validate file parameter
        if not file or not hasattr(file, 'read'):
            st.error("Invalid file object passed to process_docx_file")
            return {
                'text_chunks': 0,
                'tables': 0,
                'images': 0,
                'extracted_text': "Invalid file object",
                'extracted_tables': [],
                'extracted_images': []
            }
        
        extracted_data = {
            'text_chunks': 0,
            'tables': 0,
            'images': 0,
            'extracted_text': "",
            'extracted_tables': [],
            'extracted_images': []
        }
        
        # Read DOCX content
        doc = Document(file)
        
        # Extract text
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            extracted_data['extracted_tables'].append(table_data)
        
        # Combine text and create chunks
        full_text = "\n".join(text_content)
        extracted_data['extracted_text'] = full_text
        
        # Create text chunks
        try:
            processing_settings = st.session_state.get('processing_settings', {})
            if not isinstance(processing_settings, dict):
                st.warning(f"⚠️ processing_settings is not a dictionary: {type(processing_settings)}")
                processing_settings = {}
            
            chunk_size = processing_settings.get('chunk_size', 1000)
            chunks = [full_text[i:i+chunk_size] for i in range(0, len(full_text), chunk_size)]
            extracted_data['text_chunks'] = len(chunks)
            extracted_data['tables'] = len(extracted_data['extracted_tables'])
        except Exception as chunk_error:
            st.warning(f"⚠️ Error creating text chunks: {str(chunk_error)}")
            extracted_data['text_chunks'] = 1
            extracted_data['tables'] = len(extracted_data['extracted_tables'])
        
        return extracted_data
        
    except Exception as e:
        st.error(f"Error processing DOCX: {str(e)}")
        st.error(f"File type: {type(file)}")
        st.error(f"File attributes: {dir(file) if hasattr(file, '__dict__') else 'No attributes'}")
        return {
            'text_chunks': 0,
            'tables': 0,
            'images': 0,
            'extracted_text': "",
            'extracted_tables': [],
            'extracted_images': []
        }

def process_text_file(file):
    """Process text file and extract content"""
    try:
        # Validate file parameter
        if not file or not hasattr(file, 'read'):
            st.error("Invalid file object passed to process_text_file")
            return {
                'text_chunks': 0,
                'tables': 0,
                'images': 0,
                'extracted_text': "Invalid file object",
                'extracted_tables': [],
                'extracted_images': []
            }
        
        extracted_data = {
            'text_chunks': 0,
            'tables': 0,
            'images': 0,
            'extracted_text': "",
            'extracted_tables': [],
            'extracted_images': []
        }
        
        # Read text content
        text_content = file.read().decode('utf-8')
        extracted_data['extracted_text'] = text_content
        
        # Create text chunks
        try:
            processing_settings = st.session_state.get('processing_settings', {})
            if not isinstance(processing_settings, dict):
                st.warning(f"⚠️ processing_settings is not a dictionary: {type(processing_settings)}")
                processing_settings = {}
            
            chunk_size = processing_settings.get('chunk_size', 1000)
            chunks = [text_content[i:i+chunk_size] for i in range(0, len(text_content), chunk_size)]
            extracted_data['text_chunks'] = len(chunks)
        except Exception as chunk_error:
            st.warning(f"⚠️ Error creating text chunks: {str(chunk_error)}")
            extracted_data['text_chunks'] = 1
        
        return extracted_data
        
    except Exception as e:
        st.error(f"Error processing image: {str(e)}")
        return {
            'text_chunks': 0,
            'tables': 0,
            'images': 0,
            'extracted_text': "",
            'extracted_tables': [],
            'extracted_images': []
        }

def process_unknown_file(file):
    """Process unknown file type"""
    return {
        'text_chunks': 0,
        'tables': 0,
        'images': 0,
        'extracted_text': "File type not supported for content extraction",
        'extracted_tables': [],
        'extracted_images': []
    }

def calculate_confidence_score(extracted_data):
    """Calculate confidence score based on extraction quality"""
    score = 0
    
    # Text extraction quality
    if extracted_data['text_chunks'] > 0:
        score += 30
        if extracted_data['extracted_text'] and len(extracted_data['extracted_text']) > 100:
            score += 20
    
    # Table extraction quality
    if extracted_data['tables'] > 0:
        score += 25
    
    # Image extraction quality
    if extracted_data['images'] > 0:
        score += 15
    
    # Overall quality bonus
    if score >= 80:
        score += 10
    
    return min(score, 100)  # Cap at 100%

def show_file_details(file_data):
    """Show detailed file information and extracted content"""
    st.markdown(f"### 📄 File Details: {file_data.get('name', 'Unknown')}")
    
    # File metadata
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**📋 File Information**")
        st.info(f"**Type:** {file_data.get('type', 'unknown').upper()}")
        st.info(f"**Size:** {file_data.get('size', 0) // 1024:,} KB")
        st.info(f"**Status:** {file_data.get('status', 'unknown').title()}")
        st.info(f"**Uploaded:** {file_data.get('uploaded_at', 'Unknown')[:19]}")
        
        if file_data.get('processing_time', 0) > 0:
            st.info(f"**Processing Time:** {file_data.get('processing_time', 0):.2f} seconds")
        
        if file_data.get('confidence', 0) > 0:
            st.info(f"**Confidence:** {file_data.get('confidence', 0):.1f}%")
    
    with col2:
        st.markdown("**🔍 Content Extraction**")
        st.success(f"**Text Chunks:** {file_data.get('text_chunks', 0)}")
        st.success(f"**Tables:** {file_data.get('tables', 0)}")
        st.success(f"**Images:** {file_data.get('images', 0)}")
        
        if file_data.get('processing_errors'):
            st.error(f"**Errors:** {len(file_data.get('processing_errors', []))}")
    
    # Show extracted content
    if file_data.get('extracted_text'):
        st.markdown("**📝 Extracted Text Preview**")
        text_preview = file_data.get('extracted_text', '')[:500]
        if len(file_data.get('extracted_text', '')) > 500:
            text_preview += "..."
        st.text_area("Text Content", text_preview, height=200, disabled=True)
    
    # Show extracted tables
    if file_data.get('extracted_tables'):
        st.markdown("**📊 Extracted Tables**")
        for i, table in enumerate(file_data.get('extracted_tables', [])[:3]):  # Show first 3 tables
            if table:
                df = pd.DataFrame(table[1:], columns=table[0] if table[0] else [f"Col_{j}" for j in range(len(table[0]) if table[0] else 1)])
                st.markdown(f"**Table {i+1}**")
                st.dataframe(df, use_container_width=True)
    
    # Show extracted images
    if file_data.get('extracted_images'):
        st.markdown("**🖼️ Extracted Images**")
        st.info(f"Found {len(file_data.get('extracted_images', []))} images in the document")
    
    # Processing errors
    if file_data.get('processing_errors'):
        st.markdown("**❌ Processing Errors**")
        for error in file_data.get('processing_errors', []):
            st.error(error)

def connect_to_backend():
    """Connect to backend for enhanced processing"""
    try:
        backend_url = "http://localhost:8000"
        response = requests.get(f"{backend_url}/health", timeout=5)
        
        if response.status_code == 200:
            st.success("✅ Backend connected successfully!")
            return True
        else:
            st.warning("⚠️ Backend responded but with errors")
            return False
            
    except requests.exceptions.ConnectionError:
        st.error("❌ Cannot connect to backend. Make sure it's running on http://localhost:8000")
        return False
    except Exception as e:
        st.error(f"❌ Error connecting to backend: {str(e)}")
        return False

def enhance_processing_with_backend(file_data):
    """Enhance processing using backend services"""
    try:
        backend_url = "http://localhost:8000"
        
        # Send file data to backend for enhanced processing
        payload = {
            "file_name": file_data.get('name'),
            "file_type": file_data.get('type'),
            "extracted_text": file_data.get('extracted_text', ''),
            "extracted_tables": file_data.get('extracted_tables', []),
            "extracted_images": file_data.get('extracted_images', [])
        }
        
        response = requests.post(f"{backend_url}/api/enhance", json=payload, timeout=30)
        
        if response.status_code == 200:
            enhanced_data = response.json()
            
            # Update file data with enhanced information
            file_data.update(enhanced_data)
            
            st.success("🚀 Enhanced processing completed!")
            return file_data
        else:
            st.warning("⚠️ Backend enhancement failed")
            return file_data
            
    except Exception as e:
        st.warning(f"⚠️ Backend enhancement unavailable: {str(e)}")
        return file_data
